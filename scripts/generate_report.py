# -*- coding: utf-8 -*-
"""
Generates a ~20 page Word report (Rapport_Taches.docx) about the 5 tasks
of the Application Mallette surveillance project.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

OUTPUT = r"c:\Users\moreo\Documents\applicationmalette\Rapport_Taches.docx"
DIAG   = r"c:\Users\moreo\Documents\applicationmalette\scripts\diagrams"

def add_figure(doc, filename, caption, width_cm=16):
    path = os.path.join(DIAG, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ---------- helpers ----------
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def style_title(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def para(doc, text, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    set_paragraph_bg(p, "F2F2F2")
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    return p

def set_paragraph_bg(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], "1F4E79")
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ---------- build doc ----------
doc = Document()

# Page setup A4
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===== COVER =====
for _ in range(4):
    doc.add_paragraph()
para(doc, "RAPPORT DE PROJET", bold=True, size=28,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x4E, 0x79))
para(doc, "Application Mallette de Surveillance", bold=True, size=20,
     align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para(doc, "Système IoT de supervision multi-capteurs", italic=True, size=14,
     align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    doc.add_paragraph()
para(doc, "Tâches réalisées :", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "1. Module capteur de fumée (acquisition + MQTT)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "2. Visualisation des données capteur sur l'application", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "3. Ajout et suppression de modules", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "4. Création du MCD de la base de données", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "5. Serveur, comptes utilisateurs, sécurité, RGPD", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    doc.add_paragraph()
para(doc, f"Date : {date.today().strftime('%d/%m/%Y')}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "Technologies : Qt 6 / C++ / MQTT / MySQL / Raspberry Pi", size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
add_page_break(doc)

# ===== TABLE OF CONTENTS =====
style_title(doc, "Sommaire", level=1)
toc = [
    ("Introduction", "3"),
    ("1. Module capteur de fumée (acquisition et client MQTT)", "4"),
    ("   1.1 Présentation du capteur MQ-2", "4"),
    ("   1.2 Acquisition de la donnée côté Raspberry Pi", "5"),
    ("   1.3 Client MQTT côté application Qt", "6"),
    ("   1.4 Sécurisation MQTT par TLS", "7"),
    ("2. Visualisation de la donnée capteur sur l'application", "8"),
    ("   2.1 Architecture du widget SmokeSensorWidget", "8"),
    ("   2.2 Graphique temps réel et historique", "9"),
    ("   2.3 Seuils, alarmes et états visuels", "10"),
    ("3. Création de l'ajout et de la suppression de modules", "11"),
    ("   3.1 Concept de module dynamique", "11"),
    ("   3.2 Le ModuleManager", "12"),
    ("   3.3 Le WidgetEditor (édition par module)", "13"),
    ("4. Création du MCD de la base de données", "14"),
    ("   4.1 Analyse des besoins", "14"),
    ("   4.2 Modèle Conceptuel de Données (MCD)", "15"),
    ("   4.3 Passage au modèle logique (MLD)", "16"),
    ("5. Serveur, comptes utilisateurs, sécurité, RGPD", "17"),
    ("   5.1 Installation du serveur WAMP / MariaDB", "17"),
    ("   5.2 Création des comptes utilisateurs et RBAC", "18"),
    ("   5.3 Sécurisation de la base de données", "19"),
    ("   5.4 Conformité au RGPD", "20"),
    ("Conclusion", "21"),
]
for label, page in toc:
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5))
    p.add_run(label)
    p.add_run("\t" + page)
add_page_break(doc)

# ===== INTRODUCTION =====
style_title(doc, "Introduction", level=1)
para(doc,
     "Ce rapport présente l'ensemble des travaux réalisés dans le cadre du projet "
     "« Application Mallette de Surveillance ». L'objectif global du projet est de "
     "concevoir une mallette pédagogique et opérationnelle permettant de superviser, "
     "depuis une application bureau développée en Qt 6 / C++, un parc de capteurs "
     "embarqués sur des Raspberry Pi (température DHT22, fumée MQ-2, caméra RTSP, "
     "qualité de l'air, etc.).")
para(doc,
     "L'architecture retenue est une architecture distribuée orientée messages : "
     "les Raspberry Pi acquièrent les données et les publient sur un broker MQTT "
     "(Mosquitto), un pont MQTT vers MySQL persiste les mesures, et l'application "
     "Qt se connecte simultanément au broker (lecture temps réel) et à la base de "
     "données (lecture historique). Cette dissociation permet à l'application de "
     "rester légère, en lecture seule sur les données capteurs, et de fonctionner "
     "même hors-ligne grâce à un mode simulation.")
para(doc,
     "Le présent document détaille les cinq tâches qui ont été confiées :")
bullet(doc, "Le module capteur de fumée : acquisition de la donnée matérielle et publication via un client MQTT.")
bullet(doc, "La visualisation de cette donnée dans l'application (widget temps réel + graphique d'historique).")
bullet(doc, "L'ajout et la suppression dynamique de modules dans le tableau de bord.")
bullet(doc, "La modélisation conceptuelle (MCD) puis logique de la base de données.")
bullet(doc, "L'installation du serveur, la création des comptes utilisateurs, la sécurisation et la mise en conformité RGPD.")
para(doc,
     "Chaque tâche est présentée selon le même plan : contexte fonctionnel, choix "
     "techniques, implémentation effective dans le code, validation et points "
     "d'attention. Les extraits de code, captures schématiques (MCD / MLD) et tables "
     "récapitulatives permettent de tracer précisément les livrables.")

add_figure(doc, "01_architecture_globale.png",
           "Figure 1 — Architecture globale du système")

para(doc,
     "La figure 1 ci-dessus illustre la chaîne complète : les capteurs sont reliés "
     "aux Raspberry Pi qui publient en MQTT sécurisé sur le broker Mosquitto. "
     "Le broker alimente d'une part la base MySQL via un pont Python, d'autre part "
     "l'application Qt 6 abonnée aux topics. L'application interroge également la "
     "base pour reconstituer l'historique.")
add_page_break(doc)

# =====================================================
# TÂCHE 1 - SMOKE SENSOR MODULE
# =====================================================
style_title(doc, "1. Module capteur de fumée (acquisition et client MQTT)", level=1)

style_title(doc, "1.1 Présentation du capteur MQ-2", level=2)
para(doc,
     "Le capteur retenu pour la détection de fumée est le module MQ-2, un capteur "
     "à oxyde d'étain (SnO2) sensible aux gaz combustibles : fumée, GPL, propane, "
     "méthane, hydrogène, alcool et CO. Il dispose d'une sortie analogique (AO) "
     "proportionnelle à la concentration mesurée, et d'une sortie numérique (DO) "
     "configurée via un potentiomètre de seuil. Dans notre déploiement, c'est la "
     "sortie analogique qui est exploitée afin d'obtenir une mesure continue, "
     "exprimée ensuite en ppm (parties par million) après linéarisation.")

make_table(doc,
    ["Caractéristique", "Valeur"],
    [
        ["Référence", "MQ-2"],
        ["Tension d'alimentation", "5 V DC"],
        ["Sortie", "Analogique (AO) + Numérique (DO)"],
        ["Plage de détection", "300 – 10 000 ppm"],
        ["Temps de chauffe", "≥ 20 secondes"],
        ["Identifiant capteur (BDD)", "smoke-001"],
        ["Nœud rattaché", "rpi-001"],
        ["Seuil de warning", "50 ppm"],
        ["Seuil d'alarme", "100 ppm"],
    ],
    col_widths=[6, 9]
)

para(doc,
     "Le capteur étant analogique, il est connecté à un convertisseur analogique-"
     "numérique (ADC MCP3008) relié en SPI au Raspberry Pi rpi-001. La conversion "
     "est ensuite normalisée par le programme d'acquisition Python avant publication.")

style_title(doc, "1.2 Acquisition de la donnée côté Raspberry Pi", level=2)
para(doc,
     "Le programme d'acquisition tourne en service systemd sur le Raspberry Pi. Il "
     "réalise les étapes suivantes :")
bullet(doc, "Lecture périodique (toutes les secondes) du canal SPI relié à l'ADC.")
bullet(doc, "Conversion de la valeur brute (0–1023) en ppm via une courbe de calibration logarithmique propre au MQ-2.")
bullet(doc, "Encapsulation de la mesure dans un message JSON.")
bullet(doc, "Publication sur le topic MQTT « sensors/smoke/rpi-001 ».")
bullet(doc, "Gestion d'une reconnexion automatique au broker en cas de perte de réseau.")

para(doc, "Exemple de payload publié :", italic=True)
code_block(doc,
'{\n'
'  "sensor_id": "smoke-001",\n'
'  "node_id":   "rpi-001",\n'
'  "type":      "smoke",\n'
'  "value":     37,\n'
'  "unit":      "ppm",\n'
'  "timestamp": "2026-05-26T10:42:31Z"\n'
'}')

para(doc,
     "Le choix du JSON permet une extension simple (ajout d'attributs comme la "
     "température de chauffe ou la tension d'alimentation) sans casser les "
     "consommateurs existants. Le format est également directement réutilisé par "
     "le pont MQTT → MySQL (script mqtt_to_mysql.py).")

style_title(doc, "1.3 Client MQTT côté application Qt", level=2)
para(doc,
     "Côté application bureau, un client MQTT a été intégré dans la classe "
     "MqttClient (fichiers mqttclient.h / mqttclient.cpp). Cette classe encapsule "
     "QMqttClient (module officiel Qt MQTT) et expose un signal Qt à chaque message "
     "reçu. Le DashboardWindow s'abonne à ce signal pour redistribuer les valeurs "
     "vers les widgets capteurs concernés (dont SmokeSensorWidget).")

para(doc, "Topics auxquels l'application s'abonne :", italic=True)
bullet(doc, "sensors/smoke/# – tous les capteurs de fumée du parc")
bullet(doc, "sensors/temperature/# – capteurs DHT22")
bullet(doc, "sensors/+/+ – pattern générique pour ajout futur")

para(doc,
     "Le flux complet est donc le suivant :")
code_block(doc,
"MQ-2  --SPI-->  Raspberry Pi  --MQTT(JSON)-->  Broker Mosquitto\n"
"                                                    |\n"
"                  +---------------------------------+\n"
"                  |                                 |\n"
"           mqtt_to_mysql.py                 Application Qt (MqttClient)\n"
"                  |                                 |\n"
"             MySQL (historique)           SmokeSensorWidget (temps réel)")

style_title(doc, "1.4 Sécurisation MQTT par TLS", level=2)
para(doc,
     "Afin de protéger les échanges (mots de passe, payloads capteurs) entre les "
     "Raspberry Pi, l'application bureau et le broker, une couche TLS a été "
     "déployée. Le broker Mosquitto écoute sur le port 8883, et l'authentification "
     "mutuelle (mTLS) repose sur un certificat d'autorité (CA) commun. Côté Qt, "
     "la configuration s'appuie sur QSslConfiguration (Qt 5.15+ / Qt 6) avec un "
     "chargement explicite du CA – sans fallback non chiffré pour éviter toute "
     "régression silencieuse.")
para(doc,
     "Le signal exploité pour confirmer le canal sécurisé est encrypted() de "
     "QSslSocket, et non connected() : cette précision a permis de corriger une "
     "panne d'établissement de session SSL constatée lors des premiers essais.")

add_figure(doc, "04_sequence_smoke.png",
           "Figure 2 — Séquence d'acquisition et de diffusion d'une mesure de fumée")
add_page_break(doc)

# =====================================================
# TÂCHE 2 - VISUALISATION
# =====================================================
style_title(doc, "2. Visualisation de la donnée du capteur de fumée", level=1)

style_title(doc, "2.1 Architecture du widget SmokeSensorWidget", level=2)
para(doc,
     "Le widget de visualisation est implémenté dans les fichiers "
     "smokesensorwidget.h et smokesensorwidget.cpp. C'est un QFrame autonome qui "
     "se compose de plusieurs éléments :")
bullet(doc, "Un titre éditable (QLabel + bouton ✎ d'édition).")
bullet(doc, "Une valeur courante exprimée en ppm, affichée en grande taille.")
bullet(doc, "Un libellé d'état (Normal / Warning / Alarme) coloré.")
bullet(doc, "Un graphique d'historique sur les 30 dernières mesures.")
bullet(doc, "Un bouton de fermeture (✕) pour supprimer le module du tableau de bord.")

para(doc,
     "Le widget expose une API publique permettant deux modes de fonctionnement :")
make_table(doc,
    ["Méthode", "Rôle"],
    [
        ["setRealTimeMode(bool)", "Active la consommation MQTT plutôt que la simulation."],
        ["updateFromMqtt(int)", "Injecte une nouvelle valeur reçue du broker."],
        ["simulateStep()", "Génère une valeur factice (mode démo / hors-ligne)."],
        ["resetSensor()", "Remet à zéro la valeur courante et l'historique."],
        ["currentValue() / severity()", "Lecture de l'état pour l'agrégation au niveau dashboard."],
    ],
    col_widths=[6.5, 9]
)

style_title(doc, "2.2 Graphique temps réel et historique", level=2)
para(doc,
     "Pour visualiser l'évolution de la concentration de fumée, un graphique "
     "dynamique est intégré directement dans le widget. Les 30 dernières mesures "
     "sont conservées dans un QVector<int> (m_historyValues) et redessinées à "
     "chaque mise à jour via la méthode privée updateChart(). Lorsque le mode "
     "temps réel est inactif (broker inaccessible, mode démonstration), un QTimer "
     "déclenche périodiquement simulateStep() pour produire une courbe simulée "
     "réaliste – ce mécanisme de repli permet de démontrer l'application en "
     "salle, sans dépendance réseau.")
para(doc,
     "La logique de mise à jour suit le pseudo-code suivant :")
code_block(doc,
"void SmokeSensorWidget::updateChart() {\n"
"    while (m_historyValues.size() > 30)\n"
"        m_historyValues.removeFirst();\n"
"    repaintChartArea();         // points + ligne + grille\n"
"    drawWarningLine(50);        // seuil jaune\n"
"    drawAlarmLine(100);         // seuil rouge\n"
"}")

style_title(doc, "2.3 Seuils, alarmes et états visuels", level=2)
para(doc,
     "Trois niveaux de sévérité sont définis dans l'énumération Severity du "
     "widget :")
make_table(doc,
    ["Sévérité", "Condition", "Couleur du cadre"],
    [
        ["Normal", "valeur < 50 ppm", "Vert"],
        ["Warning", "50 ≤ valeur < 100 ppm", "Orange"],
        ["Alarm", "valeur ≥ 100 ppm", "Rouge clignotant"],
    ],
    col_widths=[3.5, 6.5, 5]
)
para(doc,
     "Les seuils sont configurables directement depuis le WidgetEditor (voir "
     "section 3.3), ce qui permet d'adapter le module à des environnements plus "
     "ou moins exigeants (laboratoire, atelier, salle serveur) sans recompilation. "
     "Les seuils par défaut (50 / 100 ppm) sont par ailleurs persistés dans la "
     "table sensors de la base de données.")
para(doc,
     "Sur le plan ergonomique, le widget est redimensionnable à la souris grâce "
     "au composant ResizableContainer / ResizeHandle. Le contenu (titre, valeur, "
     "graphique) se redimensionne proportionnellement, ce qui répond à "
     "l'exigence « Widget Content Auto-Resizing » du cahier des charges.")

add_figure(doc, "06_class_diagram.png",
           "Figure 3 — Diagramme de classes simplifié (côté application Qt)")
add_page_break(doc)

# =====================================================
# TÂCHE 3 - MODULE MANAGEMENT
# =====================================================
style_title(doc, "3. Ajout et suppression de modules sur l'application", level=1)

style_title(doc, "3.1 Concept de module dynamique", level=2)
para(doc,
     "Un « module » désigne, dans le vocabulaire de l'application, un widget "
     "capteur (fumée, température, caméra, etc.) que l'utilisateur peut placer "
     "librement sur le tableau de bord. Le choix initial reposait sur un "
     "QGridLayout figé, mais cette approche a été abandonnée au profit d'un "
     "positionnement absolu : chaque module est désormais un enfant direct de la "
     "zone centrale, repositionnable au pixel près et redimensionnable à la "
     "souris. Cette migration est consignée dans la mémoire « Qt Layout Migration "
     "from Grid to Absolute Positioning Workflow ».")

bullet(doc, "Ajout : depuis le bouton ⚙ du bandeau inférieur, l'utilisateur ouvre le ModuleManager et choisit le type de capteur à ajouter (smoke / temperature / camera / radiation).")
bullet(doc, "Suppression : un bouton ✕ est présent sur chaque widget capteur ; il déclenche la dissociation du widget de son ResizableContainer et la suppression de l'entrée correspondante dans la liste des modules actifs.")
bullet(doc, "Édition : un bouton ✎ ouvre le WidgetEditor pour modifier nom, type, seuils, unité.")
bullet(doc, "Réordonnancement : possible via le ModuleManager (déplacement haut/bas).")

style_title(doc, "3.2 Le ModuleManager", level=2)
para(doc,
     "La classe ModuleManager (fichiers modulemanager.h / modulemanager.cpp) "
     "centralise la gestion des modules. Elle expose une interface QDialog qui "
     "permet de :")
bullet(doc, "Ajouter un nouveau module en sélectionnant son type et son nom.")
bullet(doc, "Éditer un module existant (délégué au WidgetEditor).")
bullet(doc, "Supprimer un module (suppression du widget + nettoyage de la persistance éventuelle).")
bullet(doc, "Réordonner la liste pour refléter l'ordre d'affichage.")

para(doc, "Structure principale (extrait simplifié) :", italic=True)
code_block(doc,
"struct ModuleInfo {\n"
"    QString id;          // identifiant unique (UUID)\n"
"    QString name;        // nom affiché\n"
"    QString type;        // smoke | temperature | camera | ...\n"
"    int warningThreshold;\n"
"    int alarmThreshold;\n"
"    QString unit;\n"
"};\n"
"\n"
"class ModuleManager : public QDialog {\n"
"    Q_OBJECT\n"
"signals:\n"
"    void moduleAdded(const ModuleInfo &info);\n"
"    void moduleRemoved(const QString &id);\n"
"    void moduleEdited(const ModuleInfo &info);\n"
"};")

para(doc,
     "Le DashboardWindow connecte ces trois signaux à ses propres slots de "
     "création / destruction de widgets. L'ajout d'un module se traduit donc "
     "concrètement par la création d'un nouveau widget (instance "
     "SmokeSensorWidget, TemperatureWidget, CameraWidget, etc.) enveloppé dans "
     "un ResizableContainer, et placé à une position par défaut sur la zone de "
     "tableau de bord.")

style_title(doc, "3.3 Le WidgetEditor (édition par module)", level=2)
para(doc,
     "Le WidgetEditor (widgeteditor.h / widgeteditor.cpp) est un QDialog dédié "
     "à l'édition d'un module individuel. Il pré-remplit ses champs avec la "
     "configuration courante du widget et applique les changements à la "
     "validation. Les champs proposés diffèrent selon le type : un widget caméra "
     "expose une URL RTSP, tandis qu'un widget capteur expose des seuils numériques.")

make_table(doc,
    ["Champ", "Type capteurs", "Type caméra"],
    [
        ["Nom du module", "Oui", "Oui"],
        ["Type", "Oui (smoke / temp / radiation)", "Caméra"],
        ["Unité", "Oui (ppm, °C, μSv/h)", "Non"],
        ["Seuil de warning", "Oui", "Non"],
        ["Seuil d'alarme", "Oui", "Non"],
        ["URL RTSP", "Non", "Oui"],
    ],
    col_widths=[4, 5.5, 5]
)

para(doc,
     "Ce duo ModuleManager + WidgetEditor offre à l'utilisateur final une "
     "souplesse complète : il peut composer un tableau de bord adapté à son site "
     "(salle serveur, atelier, entrée extérieure) sans toucher au code. La liste "
     "des modules actifs constitue la « configuration utilisateur » qui pourra, "
     "dans une évolution ultérieure, être persistée en base via la table "
     "system_config.")

add_figure(doc, "08_activity_module.png",
           "Figure 4 — Diagramme d'activité : ajout / édition / suppression d'un module",
           width_cm=12)
add_page_break(doc)

# =====================================================
# TÂCHE 4 - MCD
# =====================================================
style_title(doc, "4. Création du MCD de la base de données", level=1)

style_title(doc, "4.1 Analyse des besoins", level=2)
para(doc,
     "Le SGBD cible est MySQL/MariaDB (intégré dans WAMP côté serveur). Les "
     "besoins fonctionnels identifiés ont structuré la modélisation :")
bullet(doc, "Authentifier des utilisateurs avec trois rôles (admin, opérateur, visiteur).")
bullet(doc, "Tracer toutes les actions sensibles (audit log) pour la traçabilité RGPD.")
bullet(doc, "Référencer les nœuds Raspberry Pi du parc (adresse IP, MAC, localisation, état en ligne).")
bullet(doc, "Référencer les capteurs rattachés à chaque nœud (type, unité, seuils).")
bullet(doc, "Historiser les mesures capteurs (sensor_data) avec horodatage et statut.")
bullet(doc, "Stocker la configuration globale du système (table system_config clé/valeur).")

style_title(doc, "4.2 Modèle Conceptuel de Données (MCD)", level=2)
para(doc,
     "Le MCD identifie six entités principales et leurs associations. Il est "
     "présenté ici sous forme textuelle (entités + associations + cardinalités) "
     "puis schématiquement.")

make_table(doc,
    ["Entité", "Attributs clés"],
    [
        ["USER", "id, username, password (hash), role, full_name, email, is_active, last_login"],
        ["AUDIT_LOG", "id, username, action, details, ip_address, timestamp"],
        ["RASPBERRY_NODE", "id, node_id, name, ip_address, mac_address, location, is_online, last_seen"],
        ["SENSOR", "id, sensor_id, name, type, unit, warning_threshold, alarm_threshold, is_active"],
        ["SENSOR_DATA", "id, value, raw_value, status, recorded_at"],
        ["SYSTEM_CONFIG", "id, config_key, config_value, description, updated_at, updated_by"],
    ],
    col_widths=[4.5, 11]
)

para(doc, "Associations :", bold=True)
bullet(doc, "HÉBERGE entre RASPBERRY_NODE (1, n) et SENSOR (1, 1) – un nœud héberge plusieurs capteurs, un capteur appartient à un seul nœud.")
bullet(doc, "PRODUIT entre SENSOR (1, n) et SENSOR_DATA (1, 1) – un capteur produit de nombreuses mesures, chaque mesure est rattachée à un seul capteur.")
bullet(doc, "EFFECTUE entre USER (1, n) et AUDIT_LOG (1, 1) – un utilisateur effectue plusieurs actions tracées.")
bullet(doc, "MAINTIENT entre USER (1, n) et SYSTEM_CONFIG (0, 1) via updated_by – un administrateur peut être l'auteur de la dernière mise à jour d'une clé de configuration.")

para(doc, "Schéma simplifié (notation textuelle) :", italic=True)
add_figure(doc, "02_mcd.png",
           "Figure 5 — Modèle Conceptuel de Données (MCD)")

style_title(doc, "4.3 Passage au modèle logique (MLD)", level=2)
para(doc,
     "Les cardinalités (1,n) – (1,1) se traduisent par l'ajout d'une clé "
     "étrangère côté entité « N ». Le MLD résultant correspond aux six tables "
     "définies dans database/surveillance_schema.sql :")

make_table(doc,
    ["Table", "Clé primaire", "Clé(s) étrangère(s)"],
    [
        ["users", "id", "—"],
        ["audit_log", "id", "username (référence logique vers users.username)"],
        ["raspberry_nodes", "id", "—"],
        ["sensors", "id", "node_id → raspberry_nodes.node_id (ON DELETE CASCADE)"],
        ["sensor_data", "id (BIGINT)", "sensor_id (référence logique vers sensors.sensor_id)"],
        ["system_config", "id", "updated_by (référence logique vers users.username)"],
    ],
    col_widths=[4, 3.5, 8]
)

add_figure(doc, "03_mld.png",
           "Figure 6 — Modèle Logique de Données (MLD) — 6 tables MySQL")

para(doc,
     "Remarque importante sur la conception : l'historique sensor_data utilise "
     "un BIGINT en clé primaire pour absorber un fort volume de mesures, ainsi "
     "que des index sur (sensor_id), (recorded_at) et (status) afin d'accélérer "
     "les requêtes de tableau de bord et de rétention. Une rétention de 90 jours "
     "est paramétrée par défaut dans system_config (clé data_retention_days), "
     "conforme à l'exigence de minimisation des données du RGPD.")
add_page_break(doc)

# =====================================================
# TÂCHE 5 - SERVER & SECURITY & RGPD
# =====================================================
style_title(doc, "5. Serveur, comptes utilisateurs, sécurité et RGPD", level=1)

style_title(doc, "5.1 Installation du serveur WAMP / MariaDB", level=2)
para(doc,
     "Le serveur cible est une station Windows hébergeant WAMP (Apache + MySQL/"
     "MariaDB + PHP). Les étapes de mise en place sont les suivantes :")
bullet(doc, "Installation de WAMP Server (version stable).")
bullet(doc, "Démarrage du service MySQL via le tray WAMP.")
bullet(doc, "Ouverture de phpMyAdmin et import du script database/surveillance_schema.sql.")
bullet(doc, "Vérification de la création de la base surveillance_db et des 6 tables associées.")
bullet(doc, "Déploiement du pont mosquitto/mqtt_to_mysql.py côté broker (200.26.16.180), exécuté en service systemd avec reconnexion automatique.")
bullet(doc, "Mise en place du driver MySQL Qt (qsqlmysql.dll) dans le dossier sqldrivers de l'application.")

para(doc,
     "L'application Qt se connecte ensuite via QSqlDatabase en pilote QMYSQL "
     "(hôte, port 3306, base surveillance_db). L'application travaille en "
     "lecture seule sur les tables sensor_data et raspberry_nodes : les "
     "écritures de mesures sont exclusivement réalisées par le pont MQTT → MySQL. "
     "Cette séparation des responsabilités limite drastiquement la surface "
     "d'attaque côté client.")

add_figure(doc, "07_deployment.png",
           "Figure 7 — Diagramme de déploiement (nœuds physiques)")

style_title(doc, "5.2 Création des comptes utilisateurs et RBAC", level=2)
para(doc,
     "Trois rôles ont été définis, et matérialisés à la fois dans le type ENUM "
     "MySQL et dans le code Qt (filtrage des actions disponibles). À "
     "l'authentification, l'application affiche un overlay de connexion intégré "
     "(et non une fenêtre modale) ; tant que la session n'est pas ouverte, "
     "l'interface est grisée et inopérante.")

make_table(doc,
    ["Rôle", "Compte par défaut", "Droits"],
    [
        ["admin",     "admin / admin123",         "Tous droits : utilisateurs, modules, configuration système, purge historique."],
        ["operator",  "operateur / operateur123", "Ajout / édition / suppression de modules, lecture historique, pas de gestion d'utilisateurs."],
        ["viewer",    "visiteur / visiteur123",   "Lecture seule des widgets et de l'historique des capteurs."],
    ],
    col_widths=[2.5, 5, 8]
)

para(doc,
     "Les mots de passe sont stockés sous forme de hachage SHA-256 dans la "
     "colonne users.password. Une évolution prévue est le passage à un hachage "
     "salé bcrypt/Argon2 pour résister aux attaques par rainbow tables. Les "
     "mots de passe par défaut ci-dessus doivent impérativement être changés "
     "à la première mise en production.")

add_figure(doc, "05_use_case_rbac.png",
           "Figure 8 — Diagramme de cas d'utilisation (RBAC à trois rôles)")

style_title(doc, "5.3 Sécurisation de la base de données", level=2)
para(doc,
     "Plusieurs mesures de sécurisation ont été mises en œuvre, couvrant "
     "réseau, base et application :")
bullet(doc, "Compte MySQL dédié à l'application avec privilèges minimaux (SELECT sur sensor_data, sensors, raspberry_nodes ; SELECT/UPDATE sur users pour la mise à jour de last_login).")
bullet(doc, "Compte distinct pour le pont MQTT → MySQL avec INSERT sur sensor_data uniquement.")
bullet(doc, "Connexions MQTT chiffrées TLS (port 8883), authentification mutuelle par certificat (mTLS).")
bullet(doc, "Mots de passe utilisateurs hachés (SHA-256, évolution vers bcrypt).")
bullet(doc, "Journalisation systématique des actions sensibles dans audit_log (connexion, ajout/suppression de module, modification de configuration).")
bullet(doc, "Contrainte FOREIGN KEY avec ON DELETE CASCADE sur sensors → raspberry_nodes pour garantir l'intégrité référentielle.")
bullet(doc, "Sauvegardes régulières via mysqldump planifiées (tâche planifiée Windows).")
bullet(doc, "Filtrage des entrées utilisateur côté Qt (QSqlQuery préparée, jamais de concaténation SQL).")

style_title(doc, "5.4 Conformité au RGPD", level=2)
para(doc,
     "Bien que l'application traite essentiellement de données techniques "
     "(mesures capteurs), elle manipule également des données personnelles : "
     "comptes utilisateurs (identifiant, nom complet, email), traces "
     "d'utilisation (adresse IP dans audit_log). Les mesures suivantes ont été "
     "mises en place pour assurer la conformité RGPD :")

make_table(doc,
    ["Principe RGPD", "Mise en œuvre"],
    [
        ["Licéité, loyauté, transparence", "Information de l'utilisateur lors de la création de compte sur les données collectées et leur finalité (supervision)."],
        ["Limitation des finalités", "Les données ne sont utilisées que pour authentification, audit et supervision capteurs."],
        ["Minimisation",                 "Seuls les champs strictement nécessaires sont stockés (pas de date de naissance, pas de numéro de téléphone)."],
        ["Exactitude",                   "L'utilisateur peut mettre à jour son profil ; les informations sont datées (updated_at)."],
        ["Limitation de conservation",   "Rétention des mesures à 90 jours (clé data_retention_days), suppression automatique au-delà."],
        ["Intégrité et confidentialité", "Hachage des mots de passe, TLS sur MQTT, accès BDD par compte à privilèges réduits, audit log."],
        ["Responsabilité (accountability)", "Journal d'audit horodaté et conservé, comptes nominatifs (pas de compte partagé)."],
        ["Droit d'accès / rectification / effacement", "Procédure manuelle administrateur : requête sur la table users + audit_log, suppression sur demande."],
    ],
    col_widths=[5, 10]
)

para(doc,
     "Côté technique, deux points méritent d'être soulignés : (1) l'IP "
     "stockée dans audit_log est une donnée à caractère personnel et doit "
     "donc être anonymisée au-delà de la durée légale de conservation ; (2) le "
     "fichier de configuration mqtt_to_mysql.py contient des identifiants "
     "MySQL et doit être protégé par permissions strictes (chmod 600) côté "
     "serveur Linux du broker, et exclu du dépôt Git via .gitignore.")
add_page_break(doc)

# ===== CONCLUSION =====
style_title(doc, "Conclusion", level=1)
para(doc,
     "Les cinq tâches confiées dans le cadre du projet Application Mallette de "
     "Surveillance ont été menées à terme et intégrées dans une chaîne "
     "fonctionnelle homogène. Le module capteur de fumée acquiert sa donnée "
     "depuis le Raspberry Pi rpi-001, publie un message JSON sécurisé sur le "
     "broker MQTT Mosquitto (TLS, port 8883), et alimente en temps réel un "
     "widget Qt dédié qui affiche la valeur courante, l'historique et l'état "
     "de sévérité.")
para(doc,
     "La capacité d'ajouter et de supprimer dynamiquement des modules depuis "
     "l'interface (ModuleManager + WidgetEditor) confère à la mallette toute "
     "sa souplesse pédagogique et opérationnelle : un même exécutable peut "
     "être déployé sur différents sites avec des configurations capteurs "
     "très différentes.")
para(doc,
     "Sur le plan données, la modélisation MCD/MLD a abouti à un schéma "
     "MySQL en six tables, déployé via le script surveillance_schema.sql sur "
     "un serveur WAMP. Les comptes utilisateurs sont gérés en RBAC à trois "
     "niveaux (admin / operator / viewer), les mots de passe sont hachés, "
     "les échanges MQTT sont chiffrés, l'audit log assure la traçabilité, et "
     "une politique de rétention de 90 jours répond à l'exigence RGPD de "
     "limitation de conservation.")
para(doc,
     "Plusieurs pistes d'amélioration ont été identifiées pour les "
     "itérations suivantes : passage à un hachage bcrypt/Argon2 avec sel, "
     "mise en place d'un système d'export RGPD automatisé (droit d'accès), "
     "ajout d'alertes email (paramètre alert_email_enabled déjà présent en "
     "base) et persistance de la configuration utilisateur du tableau de "
     "bord. L'architecture mise en place – MQTT côté événements, MySQL côté "
     "historique, Qt côté présentation – est suffisamment modulaire pour "
     "absorber ces évolutions sans refonte.")
para(doc, " ")
para(doc, "— Fin du rapport —", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
doc.save(OUTPUT)
print("OK ->", OUTPUT)
